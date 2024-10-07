from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.primitives import serialization, hashes
from cryptography.hazmat.backends import default_backend
import pandas as pd
from docx import Document
import base64

# RSA key generation
def generate_rsa_keypair():
    private_key = rsa.generate_private_key(
        public_exponent=65537,
        key_size=2048,
        backend=default_backend()
    )
    public_key = private_key.public_key()
    return private_key, public_key

# Encrypt data using RSA
def rsa_encrypt(public_key, plaintext):
    encrypted = public_key.encrypt(
        plaintext.encode('utf-8'),
        padding.OAEP(
            mgf=padding.MGF1(algorithm=hashes.SHA256()),
            algorithm=hashes.SHA256(),
            label=None
        )
    )
    return base64.b64encode(encrypted).decode('utf-8')  # Encode as Base64 to store as text

# Decrypt data using RSA
def rsa_decrypt(private_key, ciphertext):
    decrypted = private_key.decrypt(
        base64.b64decode(ciphertext),  # Decode Base64 back to bytes
        padding.OAEP(
            mgf=padding.MGF1(algorithm=hashes.SHA256()),
            algorithm=hashes.SHA256(),
            label=None
        )
    )
    return decrypted.decode('utf-8')

# Read PII from Excel
def read_pii_from_excel(excel_file):
    df = pd.read_excel(excel_file)
    return df['text'].tolist(), df['label'].tolist()

# Overwrite encrypted data in DOCX
def overwrite_encrypted_data_in_docx(docx_file, pii_data, encrypted_data):
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        for i, text in enumerate(pii_data):
            if pii_data[i] in paragraph.text:
                # Replace PII with the encrypted (Base64) data
                paragraph.text = paragraph.text.replace(text, encrypted_data[i])
    doc.save("Testing_50_encrypted.docx")
    print("Encrypted data saved in 'Testing_50_encrypted.docx'.")

# Overwrite decrypted data in DOCX
def overwrite_decrypted_data_in_docx(docx_file, pii_data, decrypted_data):
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        for i, encrypted_text in enumerate(pii_data):
            if encrypted_text in paragraph.text:
                # Replace encrypted data with decrypted text
                paragraph.text = paragraph.text.replace(encrypted_text, decrypted_data[i])
    doc.save("Testing_50_decrypted.docx")
    print("Decrypted data saved in 'Testing_80_decrypted.docx'.")

if __name__ == "__main__":
    # File paths
    EXCEL_FILE = "C:/Users/Lokghesh VAK/OneDrive/Desktop/TEST_FILES/Data_Mana.xlsx"
    DOCX_FILE = "C:/Users/Lokghesh VAK/OneDrive/Desktop/TEST_FILES/Testing_80.docx"

    # Generate RSA key pair
    private_key, public_key = generate_rsa_keypair()

    # Read PII data and labels from Excel
    pii_data, pii_identifiers = read_pii_from_excel(EXCEL_FILE)

    # Encrypt PII data
    encrypted_data = []
    for text, label in zip(pii_data, pii_identifiers):
        if label == 'PII':
            encrypted_data.append(rsa_encrypt(public_key, text))
        else:
            encrypted_data.append(text)  # Leave non-PII data unchanged

    # Overwrite encrypted data in DOCX
    overwrite_encrypted_data_in_docx(DOCX_FILE, pii_data, encrypted_data)

    # Decrypt the encrypted data
    decrypted_data = []
    for data, label in zip(encrypted_data, pii_identifiers):
        if label == 'PII':
            decrypted_data.append(rsa_decrypt(private_key, data))
        else:
            decrypted_data.append(data)  # Non-PII data remains unchanged

    # Overwrite decrypted data in another DOCX file
    overwrite_decrypted_data_in_docx("Testing_50_encrypted.docx", encrypted_data, decrypted_data)
