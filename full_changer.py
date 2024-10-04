from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.primitives import serialization, hashes
from cryptography.hazmat.backends import default_backend
import pandas as pd
from docx import Document
import base64

# RSA key generation
def generate_rsa_keypair():
    private_key = rsa.generate_private_key(
        public_exponent=65537,
        key_size=2048,
        backend=default_backend()
    )
    public_key = private_key.public_key()
    return private_key, public_key

# Encrypt data using RSA
def rsa_encrypt(public_key, plaintext):
    encrypted = public_key.encrypt(
        plaintext.encode('utf-8'),
        padding.OAEP(
            mgf=padding.MGF1(algorithm=hashes.SHA256()),
            algorithm=hashes.SHA256(),
            label=None
        )
    )
    return base64.b64encode(encrypted).decode('utf-8')  

# Decrypt data using RSA
def rsa_decrypt(private_key, ciphertext):
    decrypted = private_key.decrypt(
        base64.b64decode(ciphertext),  
        padding.OAEP(
            mgf=padding.MGF1(algorithm=hashes.SHA256()),
            algorithm=hashes.SHA256(),
            label=None
        )
    )
    return decrypted.decode('utf-8')


def read_pii_from_excel(excel_file):
    df = pd.read_excel(excel_file)
    return df['text'].tolist(), df['label'].tolist(), df['Type of PII'].tolist()  # Assuming 'Type of PII' is the third column


def save_encrypted_data_to_docx(docx_file, pii_data, encrypted_data, pii_labels):
    doc = Document()
    for text, enc_text, label in zip(pii_data, encrypted_data, pii_labels):
        if label == 'PII':
            doc.add_paragraph(enc_text)
        else:
            doc.add_paragraph(text)  
    doc.save(docx_file)
    print(f"Encrypted data saved in '{docx_file}'.")


def save_decrypted_data_to_docx(docx_file, pii_data, decrypted_data, pii_labels):
    doc = Document()
    for text, dec_text, label in zip(pii_data, decrypted_data, pii_labels):
        if label == 'PII':
            doc.add_paragraph(dec_text)
        else:
            doc.add_paragraph(text)  
    doc.save(docx_file)
    print(f"Decrypted data saved in '{docx_file}'.")

if __name__ == "__main__":
    # File paths
    EXCEL_FILE = "C:/Users/Lokghesh VAK/OneDrive/Desktop/TEST_FILES/Data_Mana.xlsx"
    ENCRYPTED_FILE = "C:/Users/Lokghesh VAK/OneDrive/Desktop/TEST_FILES/Testing_.docx"
    DECRYPTED_FILE = "C:/Users/Lokghesh VAK/OneDrive/Desktop/TEST_FILES/Testing_files_decrypted.docx"

    # Generate RSA key pair
    private_key, public_key = generate_rsa_keypair()

    # Read PII data and labels from Excel
    pii_data, pii_identifiers, pii_categories = read_pii_from_excel(EXCEL_FILE)

    # Encrypt PII data
    encrypted_data = []
    for text, label in zip(pii_data, pii_identifiers):
        if label == 'PII':
            encrypted_data.append(rsa_encrypt(public_key, text))
        else:
            encrypted_data.append(text)  
    save_encrypted_data_to_docx(ENCRYPTED_FILE, pii_data, encrypted_data, pii_identifiers)
    print("Which category would you like to decrypt?")
    print("1) Email Address")
    print("2) Credit Card Number")
    print("3) Phone Number")
    print("4) Aadhaar Number")
    choice = input("Enter the number corresponding to your choice: ")

    pii_types = {
        "1": "Email Address",
        "2": "Credit Card Number",
        "3": "Phone Number",
        "4": "Aadhaar Number"
    }

    selected_type = pii_types.get(choice)
    decrypted_data = []
    for data, label, category in zip(encrypted_data, pii_identifiers, pii_categories):
        if label == 'PII' and category == selected_type:
            decrypted_data.append(rsa_decrypt(private_key, data))
        else:
            decrypted_data.append(data)  # Leave non-selected PII and non-PII data unchanged

    save_decrypted_data_to_docx(DECRYPTED_FILE, pii_data, decrypted_data, pii_identifiers)
