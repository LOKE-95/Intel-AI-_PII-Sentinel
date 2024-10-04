from faker import Faker
from docx import Document
import random

fake = Faker()

def indian_phone_no(phone_num):
    digits = "".join(filter(str.isdigit, phone_num))
    if len(digits) >= 10:
        return str(f"+91{digits[:10]}")
    else:
        return str(f"+91{''.join([str(random.randint(0, 9)) for _ in range(10)])}")

def create_file(filename):
    doc = Document()
    no_records = 50000
    for _ in range(no_records):
        email = fake.email()
        phone_num = indian_phone_no(fake.phone_number())
        credit_card = fake.credit_card_number(card_type="mastercard")
        aadhaar_number = f"{fake.random_number(digits=4)} {fake.random_number(digits=4)} {fake.random_number(digits=4)}"
        
       
        doc.add_paragraph(email)
        doc.add_paragraph(phone_num)
        gib_num = 6
        for _ in range(gib_num):
            gib_text = fake.text(max_nb_chars=200)
            doc.add_paragraph(gib_text)
        
        doc.add_paragraph(credit_card)
        
        doc.add_paragraph(aadhaar_number)
        doc.add_paragraph() 

    doc.save(filename)  

print("Enter file number:")
numbers = int(input())        
saves_name = f"Testing_{numbers}"
doc_path = fr"C:\Users\Lokghesh VAK\OneDrive\Desktop\TEST_FILES\{saves_name}.docx"
create_file(doc_path)
print(f"Fake PII and gibberish data has been written to {saves_name}")
